#!/usr/bin/env python3
"""Extract content from the 3 MEAL documents for analysis."""

import sys
from pathlib import Path

BASE = Path(r"C:\Users\BertrandTINGUININ\.openclaw-autoclaw\workspace\.autoclaw-attachments")

# 1. DOCX: Rapport narratif
docx_path = BASE / "20260619-111639-4227eb98-ac2-Rapport_narratif_progres S&E.docx"
print("=" * 70)
print("DOC 1 : RAPPORT NARRATIF PROGRÈS S&E")
print("=" * 70)

from docx import Document
doc = Document(str(docx_path))

# Paragraphes
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        style = p.style.name if p.style else "Normal"
        if "Heading" in style:
            print(f"\n### [{style}] {p.text.strip()}")
        else:
            print(f"\n{p.text.strip()}")

# Tableaux
for ti, table in enumerate(doc.tables):
    print(f"\n--- TABLEAU {ti+1} ---")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))

print("\n\n")

# 2. XLSX: SE_
xlsx1_path = BASE / "20260619-111639-c634a81a-d01-SE_.xlsx"
print("=" * 70)
print("DOC 2 : SE_.xlsx")
print("=" * 70)

import openpyxl
wb1 = openpyxl.load_workbook(str(xlsx1_path), data_only=True)
for sheet_name in wb1.sheetnames:
    ws = wb1[sheet_name]
    print(f"\n### Feuille: {sheet_name} ({ws.max_row} lignes x {ws.max_column} colonnes)")
    for ri, row in enumerate(ws.iter_rows(values_only=True)):
        vals = [str(c) if c is not None else "" for c in row]
        if any(v.strip() for v in vals):
            print(" | ".join(vals))
        if ri > 60:
            print(f"... ({ws.max_row - ri - 1} lignes restantes)")
            break

print("\n\n")

# 3. XLSX: Corrigé MEAL
xlsx2_path = BASE / "20260619-111639-923e2cb9-4e6-Corrigé_MEAL_Analyse_Excel_Activité 3.xlsx"
print("=" * 70)
print("DOC 3 : Corrigé MEAL Analyse Excel Activité 3")
print("=" * 70)

wb2 = openpyxl.load_workbook(str(xlsx2_path), data_only=True)
for sheet_name in wb2.sheetnames:
    ws = wb2[sheet_name]
    print(f"\n### Feuille: {sheet_name} ({ws.max_row} lignes x {ws.max_column} colonnes)")
    for ri, row in enumerate(ws.iter_rows(values_only=True)):
        vals = [str(c) if c is not None else "" for c in row]
        if any(v.strip() for v in vals):
            print(" | ".join(vals))
        if ri > 60:
            print(f"... ({ws.max_row - ri - 1} lignes restantes)")
            break
